"""
SmartDocQA - Question Answering Service
========================================
Retrieval pipeline:

    User question
        │
        ▼
  ┌──────────────────────┐
  │  FAISS  Vector Search│  ← BGE semantic embeddings
  │  + BM25  Keyword Sz  │  ← jieba tokenization
  └──────────┬───────────┘
             │  RRF Fusion (Reciprocal Rank Fusion)
             ▼
  ┌──────────────────────┐
  │  Cross-Encoder Rank  │  ← BGE reranker re-sorts
  └──────────┬───────────┘
             │  Top-K chunks
             ▼
  ┌──────────────────────┐
  │  LLM Generation      │  ← context + history → answer + sources
  └──────────────────────┘
"""

import csv
import os
import pickle
import shutil
import time
import uuid
from collections import OrderedDict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ── HuggingFace mirror MUST be set BEFORE HF-lib imports ──
# Read HF_ENDPOINT from .env manually so it's in os.environ
# before huggingface_hub / transformers cache their default endpoint.
_env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".env")
if os.path.isfile(_env_path):
    with open(_env_path, encoding="utf-8") as _f:
        for _line in _f:
            _line = _line.strip()
            if _line.startswith("HF_ENDPOINT="):
                _val = _line.split("=", 1)[1].strip().strip("\"'")
                if _val:
                    os.environ["HF_ENDPOINT"] = _val
                break

from openai import OpenAI
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import get_settings
from app.services.structured_store import load_tables, save_tables, search_tables

settings = get_settings()

# ═══════════════════════════════════════════════════════════
#  SYSTEM PROMPT  — improved with source-citation rules
# ═══════════════════════════════════════════════════════════

DOMAIN_PROMPTS = {
    "enterprise": (
        "你是一个专业的企业内部知识助手，专门回答员工关于公司制度、技术文档、工作流程的问题。\n\n"
        "回答规则：\n"
        "1. 【结论优先】先用一句话直接回答问题核心。\n"
        "2. 【依据详实】从提供的文档内容中提取具体条款、数据或步骤作为依据。如果涉及制度规定，引用相关条款号。如果涉及流程，列出具体步骤。\n"
        "3. 【引用来源】在每个依据后使用[¹][²][³]等上标标注来自哪个文档，引用标记必须在原文中出现过。\n"
        "4. 【专业规范】用词正式、条理清晰，符合企业文档的表述风格。\n"
        "5. 【实事求是】如果文档中没有相关信息，如实告知，不要编造。\n"
        "6. 【表格数据】如果提供了表格数据（如季度报表、统计表），优先从表格中提取精确数字回答，且要标注对应的行和列信息以表明数据来源。\n"
        "7. 【趋势分析】如果提供多个时期的表格数据，自动做趋势比较，用↑/↓/→标记变化方向，并提示异常值。\n"
        "8. 【多文档对比】如果提供了多份文档的内容(如多份季度报表、多个版本方案、多个合同)，且问题涉及对比/差异/趋势看变化：\n"
        "   - 先按文档分别提取关键数据或条款\n"
        "   - 再用对比表格逐项列出各文档的异同，每行一个对比维度，每列一个文档\n"
        "   - 最后给出综合结论，明确指出各文档之间的关键差异和共性\n\n"
        "企业文档类型说明：\n"
        "- 员工手册/公司制度：涉及考勤、薪酬、绩效、休假、报销等规定\n"
        "- 技术文档：API文档、架构设计、运维手册、开发规范\n"
        "- 流程文件：SOP、审批流程、项目管理规范\n"
        "- 报表数据：季度报表、年度报表、统计表格、数据分析报告\n\n"
        "请以「结论→依据→引用」的结构组织回答。如果问题涉及数值查询或对比，优先使用表格格式输出结果。"
    ),
    "research": (
        "你是一个专业的科研文献分析助手，帮助研究人员深度分析学术论文、实验数据和科研资料。\n\n"
        "回答规则：\n"
        "1. 【核心结论】先用一句话概括研究发现或答案。\n"
        "2. 【证据支撑】引用文献中的具体数据、实验结果、方法论，用上标标注来源[¹][²][³]。\n"
        "3. 【学术严谨】保持客观、准确的学术表述，严格区分事实陈述和作者推论。\n"
        "4. 【交叉引用】如果多篇文献涉及同一问题，指出它们的共识或分歧。\n"
        "5. 【批判性分析】不仅总结文献内容，还要评估研究质量：指出实验设计的潜在缺陷、样本量是否充足、结论是否过度泛化、方法局限性、结果可复现性。使用「⚠️ 注意」标记潜在问题。\n"
        "6. 【置信度标注】对关键信息标注可信度等级：\n"
        "   - 〖高置信度〗：文献中有明确数据或实验验证的结论\n"
        "   - 〖中等置信度〗：作者提出但未充分验证的观点，或间接证据支持的结论\n"
        "   - 〖低置信度〗：推测性结论、未经实验验证的假设、或仅引用他人观点的讨论\n"
        "7. 【溯源追问】每次回答结束后，主动提出2-3个深入追问建议，引导用户继续深挖：「🔍 进一步追问：...」，帮助用户从结论追溯到实验细节。\n"
        "8. 【格式兼容】引用标注支持多种学术规范：APA格式用(Peng et al., 2023)、MLA格式用(Peng 23)、GB/T 7714格式用[1]。在来源列表中统一标注引用格式。\n"
        "9. 【跨语言分析】支持中英文混合文献分析：回答时可引用中文和英文文献，中英文术语混用时附上原文术语。对翻译内容标注「[译]」标记。\n\n"
        "科研文档类型说明：\n"
        "- 学术论文：期刊论文、会议论文、预印本\n"
        "- 研究报告：技术报告、实验报告、白皮书\n"
        "- 数据资料：数据集、统计分析、实验结果\n\n"
        "回答结构模板：\n"
        "结论：一句话概括核心发现（〖置信度〗）\n\n"
        "依据：\n"
        "- 具体证据1...〖高置信度〗[¹]\n"
        "- 具体证据2...〖中等置信度〗[²][³]\n\n"
        "⚠️ 注意：潜在缺陷或方法局限性〖中等置信度〗\n\n"
        "引用：\n"
        "[¹] 标准引用格式（APA/MLA/GB/T 7714）\n"
        "[²] ...\n\n"
        "🔍 进一步追问：\n"
        "1. 追问建议1？\n"
        "2. 追问建议2？\n"
        "3. 追问建议3？"
    ),
    "legal": (
        "你是一个专业的法律助手，帮助分析法律文件、合同条款和法规条文。\n\n"
        "回答规则：\n"
        "1. 【结论先行】先用一句话总结法律观点或分析结论。\n"
        "2. 【条文依据】引用具体的法条、合同条款、判例作为依据，标注来源[¹][²][³]。\n"
        "3. 【审慎严谨】用词精确，区分法律义务与建议，不提供确凿的法律意见。\n"
        "4. 【风险提示】如果存在法律风险点，明确指出。\n"
        "5. 【实事求是】如果文档中没有相关依据，如实告知。\n\n"
        "法律文档类型说明：\n"
        "- 合同协议：合作协议、保密协议、服务合同\n"
        "- 法规条文：法律法规、行业规范、政策文件\n"
        "- 案例资料：判决书、仲裁裁决、案例分析\n\n"
        "请以「结论→条文依据→引用」的结构组织回答。"
    ),
}

def get_domain_prompt(domain: str) -> str:
    return DOMAIN_PROMPTS.get(domain, DOMAIN_PROMPTS["enterprise"])

# ═══════════════════════════════════════════════════════════
#  SINGLETON — LLM client
# ═══════════════════════════════════════════════════════════

_llm_client: Optional[OpenAI] = None


def get_llm_client() -> OpenAI:
    """Module-level singleton — create once, reuse across all requests."""
    global _llm_client
    if _llm_client is None:
        print(f"[llm] Creating OpenAI client → {settings.openai_base_url} model={settings.openai_model}")
        _llm_client = OpenAI(
            api_key=settings.openai_api_key,
            base_url=settings.openai_base_url,
            timeout=settings.openai_timeout,
            max_retries=1,
        )
    return _llm_client


# ═══════════════════════════════════════════════════════════
#  SINGLETON EMBEDDINGS
# ═══════════════════════════════════════════════════════════

_embeddings_instance: Optional[HuggingFaceEmbeddings] = None


def get_embeddings() -> HuggingFaceEmbeddings:
    """Lazy singleton — load once, reuse across all queries."""
    global _embeddings_instance
    if _embeddings_instance is None:
        print(f"[embed] Loading {settings.local_embedding_model} ... (first time, may download)")
        _embeddings_instance = HuggingFaceEmbeddings(
            model_name=settings.local_embedding_model,
            model_kwargs={"device": "cpu"},
            encode_kwargs={
                "normalize_embeddings": True,
                "batch_size": 8,
            },
        )
        print("[embed] Model loaded OK")
    return _embeddings_instance


# ═══════════════════════════════════════════════════════════
#  SINGLETON RERANKER
# ═══════════════════════════════════════════════════════════

_reranker_instance = None  # type: Optional[object]


def get_reranker():
    """Lazy singleton cross-encoder reranker."""
    global _reranker_instance
    if _reranker_instance is None:
        try:
            from sentence_transformers import CrossEncoder
        except ImportError:
            raise RuntimeError(
                "sentence-transformers is required for reranking. "
                "Run: pip install sentence-transformers"
            )
        print(f"[rerank] Loading {settings.reranker_model} ... (first time, may download)")
        _reranker_instance = CrossEncoder(
            settings.reranker_model,
            max_length=512,
            device="cpu",
        )
        print("[rerank] Model loaded OK")
    return _reranker_instance


# ═══════════════════════════════════════════════════════════
#  BM25 KEYWORD RETRIEVER  (for hybrid search)
# ═══════════════════════════════════════════════════════════


def _tokenize(text: str) -> List[str]:
    """Chinese-aware tokenizer: uses jieba for CJK, space split for ASCII."""
    try:
        import jieba
    except ImportError:
        raise RuntimeError("jieba is required for BM25. Run: pip install jieba")
    tokens: List[str] = []
    for token in jieba.cut(text):
        token = token.strip()
        if token:
            tokens.append(token)
    return tokens


class BM25Retriever:
    """BM25 keyword retriever built over document chunks."""

    def __init__(self, documents: List[Document]):
        try:
            from rank_bm25 import BM25Okapi
        except ImportError:
            raise RuntimeError("rank-bm25 is required. Run: pip install rank-bm25")
        self._documents = documents
        self._corpus = [_tokenize(doc.page_content) for doc in documents]
        self._bm25 = BM25Okapi(self._corpus)

    def search(
        self, query: str, k: int = 10
    ) -> List[Tuple[Document, float]]:
        tokenized = _tokenize(query)
        scores = self._bm25.get_scores(tokenized)
        # Normalise to [0, 1]
        max_s = max(scores) if max(scores) > 0 else 1.0
        normalised = [s / max_s for s in scores]
        results = list(zip(self._documents, normalised))
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:k]


# ═══════════════════════════════════════════════════════════
#  RRF FUSION  (Reciprocal Rank Fusion)
# ═══════════════════════════════════════════════════════════


def _rrf_fusion(
    faiss_results: List[Tuple[Document, float]],
    bm25_results: List[Tuple[Document, float]],
    k: int = 60,
) -> List[Tuple[Document, float]]:
    """
    Merge FAISS (semantic) and BM25 (keyword) retrieval results.

    RRF formula: score(d) = Σ 1/(k + rank_i(d))

    Deduplication key: (page_content[:80], metadata["source"])
    """
    _id_of = {}  # doc -> compact id
    _doc_of = {}  # compact id -> doc
    scores: Dict[str, float] = {}

    for rank, (doc, _) in enumerate(faiss_results):
        key = _make_key(doc)
        if key not in _id_of:
            _id_of[key] = key
            _doc_of[key] = doc
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)

    for rank, (doc, _) in enumerate(bm25_results):
        key = _make_key(doc)
        if key not in _id_of:
            _id_of[key] = key
            _doc_of[key] = doc
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)

    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [(_doc_of[key], score) for key, score in ranked]


def _make_key(doc: Document) -> str:
    """Compact identity key for deduplication."""
    return f"{doc.page_content[:80]}|{doc.metadata.get('source','')}|{doc.metadata.get('page','')}|{doc.metadata.get('chunk','')}"


# ── Query type detection ────────────────────────────────

_NUMERIC_PATTERNS = [
    "多少", "几", "占比", "比例", "率", "总额", "总计", "平均",
    "最高", "最低", "最大", "最小", "排名", "排行", "前几",
    "增长", "下降", "同比", "环比", "涨幅", "跌幅",
    "数字", "数据", "统计", "数字是多少",
]

_COMPARISON_PATTERNS = [
    # Single-document perspective shifts
    "对比", "比较", "差异", "区别", "不同",
    # Cross-period
    "趋势", "变化", "走势", "波动", "变动",
    # Cross-document
    "各份", "各文档", "各文件", "所有文档", "全部文档",
    # Explicit separators (handled in function body)
]


def detect_query_type(question: str) -> str:
    """
    Classify a user question as ``"comparison"``, ``"numeric"`` or ``"semantic"``.

    ``"comparison"`` — questions asking to compare/contrast across multiple
                       documents, periods, or versions (e.g. "对比三份报表",
                       "Q1和Q2的趋势有什么不同").
    ``"numeric"``    — questions that ask for specific numbers, rankings,
                       or table-contained data.
    ``"semantic"``   — conceptual / explanatory questions answered by
                       general text retrieval.
    """
    q = question.strip().lower()

    # 1) Check comparison — do this first so "对比Q1和Q2的利润" stays comparison
    for pat in _COMPARISON_PATTERNS:
        if pat in q:
            return "comparison"

    # Explicit cross-entity separators that often indicate comparison
    if any(sep in q for sep in [" vs ", " versus ", " 和 ", "与"]):
        # "Q1和Q2的利润" is comparison; "spring和springboot区别" is also comparison
        return "comparison"

    # 2) Then numeric
    for pat in _NUMERIC_PATTERNS:
        if pat in q:
            return "numeric"

    # 3) Check for comparison between two entities (words like 对比/区别/差异 already
    #    caught above; here catch pairs like "Q1, Q2, Q3")
    if _looks_like_multi_doc_comparison(q):
        return "comparison"

    return "semantic"


def _looks_like_multi_doc_comparison(q: str) -> bool:
    """Heuristic: question mentions 2+ named periods / versions / documents."""
    # Count periods (Q1/Q2/Q3, 一季度/二季度, 1月/2月, etc.)
    period_count = 0
    for kw in ["q1", "q2", "q3", "q4", "一季度", "二季度", "三季度", "四季度",
               "1月", "2月", "3月", "4月", "5月", "6月",
               "7月", "8月", "9月", "10月", "11月", "12月",
               "版本", "v1", "v2", "v3", "第一版", "第二版"]:
        if kw in q:
            period_count += 1
    return period_count >= 2


# ═══════════════════════════════════════════════════════════
#  CROSS-ENCODER RERANK
# ═══════════════════════════════════════════════════════════


def _rerank(
    query: str,
    documents: List[Tuple[Document, float]],
    top_k: int = 4,
) -> List[Document]:
    """Re-sort candidates with a cross-encoder for precision."""
    if not documents:
        return []
    docs = [doc for doc, _ in documents]
    pairs = [[query, doc.page_content] for doc in docs]
    reranker = get_reranker()
    scores = reranker.predict(
        pairs,
        show_progress_bar=False,
        batch_size=min(8, len(docs)),
    )
    ranked = sorted(zip(docs, scores), key=lambda x: x[1], reverse=True)
    return [doc for doc, _ in ranked[:top_k]]


# ═══════════════════════════════════════════════════════════
#  CONVERSATION MEMORY  (multi-turn)
# ═══════════════════════════════════════════════════════════


@dataclass
class ConversationTurn:
    role: str          # "user" | "assistant"
    content: str
    sources: List[str] = field(default_factory=list)
    timestamp: float = field(default_factory=time.time)


class ConversationMemory:
    """In-memory multi-turn conversation store with TTL eviction."""

    def __init__(self, ttl_seconds: int = 3600):
        self._store: Dict[str, List[ConversationTurn]] = {}
        self._ttl = ttl_seconds

    def _evict(self, conv_id: str) -> None:
        now = time.time()
        if conv_id in self._store:
            self._store[conv_id] = [
                t for t in self._store[conv_id]
                if now - t.timestamp < self._ttl
            ]

    def get_history(
        self, conversation_id: str, max_turns: int = 6
    ) -> List[ConversationTurn]:
        self._evict(conversation_id)
        return self._store.get(conversation_id, [])[-max_turns:]

    def add_turn(self, conversation_id: str, turn: ConversationTurn) -> None:
        if conversation_id not in self._store:
            self._store[conversation_id] = []
        self._store[conversation_id].append(turn)
        self._evict(conversation_id)

    def clear(self, conversation_id: str) -> None:
        self._store.pop(conversation_id, None)


conversation_memory = ConversationMemory(
    ttl_seconds=settings.conversation_ttl_seconds
)


# ═══════════════════════════════════════════════════════════
#  DOCUMENT LOADER
# ═══════════════════════════════════════════════════════════


def _load_document(file_path: str) -> List[Document]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "page": i + 1},
                ))
        return docs

    elif ext == ".csv":
        with open(path, encoding="utf-8") as f:
            reader = csv.DictReader(f)
            rows = list(reader)
        if not rows:
            raise ValueError("CSV file is empty")
        docs = []
        for i, row in enumerate(rows):
            text = " | ".join(f"{k}: {v}" for k, v in row.items())
            docs.append(Document(
                page_content=text,
                metadata={"source": path.name, "row": i + 1},
            ))
        return docs

    elif ext in (".txt", ".md"):
        with open(path, encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": path.name})]

    elif ext == ".docx":
        from docx import Document as DocxDocument
        docx_doc = DocxDocument(str(path))
        text = "\n".join(
            para.text for para in docx_doc.paragraphs if para.text.strip()
        )
        return [Document(page_content=text, metadata={"source": path.name})]

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ═══════════════════════════════════════════════════════════
#  DOCUMENT SPLITTER
# ═══════════════════════════════════════════════════════════


def _split_document(documents: List[Document]) -> List[Document]:
    """Split pages into semantic chunks using RecursiveCharacterTextSplitter."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", ".", "!", "?", " "],
        keep_separator=True,
    )
    chunks = text_splitter.split_documents(documents)
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk"] = i
    return chunks


def load_and_split_document(file_path: str) -> List[Document]:
    """Public API — load + split a document file."""
    raw_docs = _load_document(file_path)
    return _split_document(raw_docs)


# ═══════════════════════════════════════════════════════════
#  QA SERVICE  (main class)
# ═══════════════════════════════════════════════════════════


class QAService:
    """
    Smart document Q&A service.

    Pipeline:  FAISS + BM25 → RRF fusion → Cross-encoder rerank → LLM
    Supports:  multi-turn conversation, source citation with chunk details.

    All heavy resources (LLM client, embeddings, reranker) are module-level
    singletons — creating a QAService instance is cheap.
    """

    # ------- Vector store CRUD -----------------------------------

    @staticmethod
    def create_vector_store(
        chunks: List[Document], vector_store_name: str
    ) -> str:
        """Create FAISS index and persist chunks for BM25 reconstruction."""
        vector_store_dir = Path(settings.vector_store_path) / vector_store_name
        if vector_store_dir.exists():
            shutil.rmtree(vector_store_dir)
        vector_store_dir.mkdir(parents=True, exist_ok=True)

        embeddings = get_embeddings()
        vector_store = FAISS.from_documents(chunks, embeddings)
        vector_store.save_local(str(vector_store_dir))

        # Verify
        idx_path = vector_store_dir / "index.faiss"
        if not idx_path.exists():
            raise RuntimeError(
                f"FAISS index not saved to {idx_path}. "
                f"Directory: {os.listdir(str(vector_store_dir))}"
            )

        # Persist chunks for BM25 rebuild
        chunks_path = vector_store_dir / "chunks.pkl"
        with open(chunks_path, "wb") as f:
            pickle.dump(chunks, f)

        # Persist version marker
        version_path = vector_store_dir / "version.txt"
        version_path.write_text(
            f"model={settings.local_embedding_model}\n", encoding="utf-8"
        )

        print(f"[vector] Created '{vector_store_name}' ({len(chunks)} chunks)")
        return str(vector_store_dir)

    @staticmethod
    def load_vector_store(vector_store_name: str):
        """Load FAISS index. Returns None if missing."""
        vector_store_dir = Path(settings.vector_store_path) / vector_store_name
        idx_path = vector_store_dir / "index.faiss"
        if not idx_path.exists():
            return None
        embeddings = get_embeddings()
        try:
            return FAISS.load_local(
                str(vector_store_dir),
                embeddings,
                allow_dangerous_deserialization=True,
            )
        except Exception as e:
            import traceback
            err = traceback.format_exc()
            # Dimension mismatch → old vector store, needs reprocessing
            raise RuntimeError(
                f"无法加载向量库 '{vector_store_name}'，可能是旧版本创建的。"
                f"请重新处理文档。错误: {e}"
            )

    @classmethod
    def _load_chunks(cls, vector_store_name: str) -> List[Document]:
        chunks_path = (
            Path(settings.vector_store_path) / vector_store_name / "chunks.pkl"
        )
        if not chunks_path.exists():
            return []
        with open(chunks_path, "rb") as f:
            return pickle.load(f)

    @staticmethod
    def delete_vector_store(vector_store_name: str) -> None:
        vector_store_dir = Path(settings.vector_store_path) / vector_store_name
        if vector_store_dir.exists():
            shutil.rmtree(vector_store_dir)

    @staticmethod
    def extract_and_store_tables(file_path: str, vector_store_name: str) -> None:
        """
        Extract tables from a PDF and persist them alongside the FAISS index.
        Safe to call on non-PDF files — no-op for .txt/.md/.csv/.docx.
        """
        ext = Path(file_path).suffix.lower()
        if ext != ".pdf":
            return
        try:
            from app.services.table_parser import extract_tables_from_pdf
            tables = extract_tables_from_pdf(file_path)
            if tables:
                store_dir = Path(settings.vector_store_path) / vector_store_name
                store_dir.mkdir(parents=True, exist_ok=True)
                save_tables(tables, str(store_dir))
        except Exception as exc:
            logger.warning(f"[qa] Table extraction skipped for {file_path}: {exc}")

    # ------- LLM chat wrapper ------------------------------------

    @staticmethod
    def llm_chat(system: str, user: str) -> str:
        client = get_llm_client()
        resp = client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.3,
        )
        return resp.choices[0].message.content

    @staticmethod
    def llm_chat_stream_generator(system: str, user: str):
        """Yield text chunks for SSE streaming."""
        client = get_llm_client()
        stream = client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.3,
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta
            if delta.content:
                yield delta.content

    def _format_answer(self, text: str) -> str:
        """Post-process LLM answer: strip markdown, normalize readability.

        Produces plain-text structured with blank lines between sections
        and each list item on its own line.
        """
        import re

        # 1 ▸ Remove markdown formatting markers
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # **bold** → bold
        text = re.sub(r'__(.+?)__', r'\1', text)        # __underline__
        text = re.sub(r'^#{1,6}\s', '', text, flags=re.MULTILINE)  # strip headers

        # 2 ▸ Ensure ordered-list items each start their own line
        #    e.g. "1.xxx，2.yyy" or "1.xxx, 2.yyy" → "1.xxx\n2.yyy"
        text = re.sub(
            r'(\d{1,2}[\.\)、]\s*\S.*?)\s*[，,]+\s*(?=\d{1,2}[\.\)、]\s)',
            r'\1\n',
            text
        )

        # 3 ▸ Ensure dash/bullet items each start their own line
        #    e.g. "- xxx，- yyy" or "- xxx, - yyy" → "- xxx\n- yyy"
        text = re.sub(
            r'([\-·•]\s*\S.*?)\s*[，,]+\s*(?=[\-·•]\s)',
            r'\1\n',
            text
        )

        # 4 ▸ Insert blank line before section labels (结论/依据/引用/分析/注意)
        for label in ['结论', '依据', '引用', '分析', '注意', '补充', '说明', '总结', '拓展', '备注']:
            text = re.sub(
                rf'(?<!\n)(?<!^)({label}[：:]\s*)',
                r'\n\n\1',
                text
            )

        # 5 ▸ Normalize excessive blank lines (3+ → 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 6 ▸ Trim leading/trailing whitespace
        text = text.strip()

        return text

    # ------- Core Q&A --------------------------------------------

    def ask_question(
        self,
        vector_store_name: str,
        question: str,
        conversation_id: str = "",
        top_k: int = 4,
        use_hybrid: bool = True,
        use_rerank: bool = True,
    ) -> dict:
        """
        Full retrieval-augmented Q&A pipeline.

        Parameters
        ----------
        vector_store_name : str   — document ID used as vector store key
        question          : str   — user's question
        conversation_id   : str   — empty = new conversation; non-empty = continue
        top_k             : int   — number of chunks to feed the LLM
        use_hybrid        : bool  — enable BM25 + FAISS hybrid retrieval
        use_rerank        : bool  — enable cross-encoder re-ranking
        """
        # 1 ▸ Load vector store
        vector_store = self.load_vector_store(vector_store_name)
        if vector_store is None:
            raise ValueError(f"Vector store not found: {vector_store_name}")

        # 2 ▸ Retrieve — FAISS (always) + BM25 (optional)
        candidates: List[Tuple[Document, float]]
        if use_hybrid:
            all_chunks = self._load_chunks(vector_store_name)
            if all_chunks:
                bm25 = BM25Retriever(all_chunks)
                bm25_results = bm25.search(question, k=top_k * 5)
            else:
                bm25_results = []
            faiss_results = vector_store.similarity_search_with_score(
                question, k=top_k * 5
            )
            candidates = _rrf_fusion(faiss_results, bm25_results, k=60)
        else:
            candidates = vector_store.similarity_search_with_score(
                question, k=top_k * 5
            )

        if not candidates:
            raise ValueError("No relevant document passages found")

        # 3 ▸ Rerank  (optional, adds precision)
        if use_rerank and len(candidates) > top_k:
            pool = candidates[: min(top_k * 3, len(candidates))]
            final_docs = _rerank(question, pool, top_k=top_k)
        else:
            final_docs = [doc for doc, _ in candidates[:top_k]]

        # 4 ▸ Build context & source details
        context_parts: List[str] = []
        source_details: OrderedDict[str, list] = OrderedDict()

        for i, doc in enumerate(final_docs):
            source = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "")
            chunk_idx = doc.metadata.get("chunk", i)

            label = f"[片段 {i + 1}]"
            if page:
                label += f"  来源: {source} · 第 {page} 页"
            else:
                label += f"  来源: {source}"
            context_parts.append(f"{label}\n{doc.page_content}")

            if source not in source_details:
                source_details[source] = []
            source_details[source].append({
                "chunk_index": i + 1,
                "page": page,
                "snippet": (
                    doc.page_content[:200] + "..."
                    if len(doc.page_content) > 200
                    else doc.page_content
                ),
            })

        context = "\n\n" + "─" * 60 + "\n\n".join(context_parts) + "\n\n" + "─" * 60

        # 5 ▸ Build user message (history + context + question)
        history_text = ""
        if conversation_id:
            history = conversation_memory.get_history(conversation_id, max_turns=6)
            if history:
                history_parts = []
                for turn in history:
                    role_label = "👤 用户" if turn.role == "user" else "🤖 助手"
                    history_parts.append(f"{role_label}: {turn.content}")
                history_text = (
                    "## 历史对话\n\n"
                    + "\n\n".join(history_parts)
                    + "\n\n" + "─" * 60 + "\n\n"
                )

        user_message = (
            f"{history_text}"
            f"## 文档片段 (共 {len(final_docs)} 段)\n"
            f"{context}\n\n"
            f"## 当前问题\n\n"
            f"{question}"
        )

        # 6 ▸ LLM generation
        answer = self.llm_chat(get_domain_prompt("enterprise"), user_message)

        # 6b ▸ Post-process answer: strip markdown, normalize formatting
        answer = self._format_answer(answer)

        # 7 ▸ Persist conversation
        convo_id = conversation_id or str(uuid.uuid4())
        conversation_memory.add_turn(
            convo_id,
            ConversationTurn(role="user", content=question),
        )
        conversation_memory.add_turn(
            convo_id,
            ConversationTurn(
                role="assistant",
                content=answer,
                sources=[s for s in source_details],
            ),
        )

        # 8 ▸ Build response
        retrieve_method = (
            "hybrid + rerank" if use_hybrid and use_rerank
            else "hybrid" if use_hybrid
            else "vector + rerank" if use_rerank
            else "vector"
        )

        return {
            "question": question,
            "conversation_id": convo_id,
            "answer": answer,
            "sources": list(source_details.keys()),
            "source_details": [
                {"source": s, "chunks": d}
                for s, d in source_details.items()
            ],
            "source_count": len(final_docs),
            "retrieval_method": retrieve_method,
        }

    # ------- Knowledge Base Q&A ----------------------------------

    def ask_question_by_kb(
        self,
        kb_id: str,
        question: str,
        conversation_id: str = "",
        top_k: int = 4,
        use_hybrid: bool = True,
        use_rerank: bool = True,
    ) -> dict:
        """Retrieve from a knowledge base (multi-document FAISS index)."""
        kb_store_name = f"kb_{kb_id}"

        # Look up KB domain
        _domain = "enterprise"
        try:
            from app.db.database import SessionLocal
            from app.models.document import KnowledgeBase
            _db = SessionLocal()
            try:
                _kb = _db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
                if _kb and _kb.domain:
                    _domain = _kb.domain
            finally:
                _db.close()
        except Exception:
            pass

        # ── 1) Classify query & load tables (if any) ──────────────
        qtype = detect_query_type(question)
        store_dir = str(Path(settings.vector_store_path) / kb_store_name)
        tables = load_tables(store_dir) if Path(store_dir).exists() else []
        table_context = ""

        if qtype == "numeric" and tables:
            matched = search_tables(tables, question)
            if matched:
                # Format matched tables as markdown for the LLM
                table_parts = []
                for i, tbl in enumerate(matched[:3]):  # top-3 tables
                    lines = [f"[表格 {i + 1}]"]
                    if tbl.get("page"):
                        lines[0] += f"  (第 {tbl['page']} 页)"
                    lines.append("")

                    # Header row
                    headers = tbl.get("headers", [])
                    if headers:
                        lines.append("| " + " | ".join(headers) + " |")
                        lines.append("|" + "|".join(["---"] * len(headers)) + "|")

                    # Data rows
                    for row in tbl.get("rows", []):
                        lines.append("| " + " | ".join(row) + " |")

                    table_parts.append("\n".join(lines))

                table_context = (
                    "\n\n## 文档中的表格数据\n\n"
                    + "\n\n".join(table_parts)
                    + "\n\n注意：表格数据优先于纯文本段落，回答数值问题时请优先从表格中提取精确数据。"
                )

        # ── 2) Text retrieval (FAISS + BM25 → RRF → Rerank) ──────
        vector_store = self.load_vector_store(kb_store_name)

        if vector_store is None:
            from app.db.database import SessionLocal
            from app.services.knowledge_base_service import KnowledgeBaseService
            db = SessionLocal()
            try:
                KnowledgeBaseService.rebuild_kb_index(db, kb_id)
            finally:
                db.close()
            vector_store = self.load_vector_store(kb_store_name)
            if vector_store is None:
                raise ValueError(f"Knowledge base index not found: {kb_id}")

        candidates = None

        if use_hybrid:
            all_chunks = self._load_chunks(kb_store_name)
            if all_chunks:
                bm25 = BM25Retriever(all_chunks)
                bm25_results = bm25.search(question, k=top_k * 5)
            else:
                bm25_results = []

            faiss_results = vector_store.similarity_search_with_score(
                question, k=top_k * 5
            )
            candidates = _rrf_fusion(faiss_results, bm25_results, k=60)

        if candidates is None:
            candidates = vector_store.similarity_search_with_score(
                question, k=top_k * 5
            )

        if not candidates:
            raise ValueError("No relevant document passages found")

        if use_rerank and len(candidates) > top_k:
            pool = candidates[: min(top_k * 3, len(candidates))]
            final_docs = _rerank(question, pool, top_k=top_k)
        else:
            final_docs = [doc for doc, _ in candidates[:top_k]]
        context_parts = []
        source_details = __import__('collections').OrderedDict()

        for i, doc in enumerate(final_docs):
            source = doc.metadata.get("source", "Unknown")
            page = str(doc.metadata.get("page") or "")
            chunk_idx = doc.metadata.get("chunk", i)
            doc_name = doc.metadata.get("document_name", source)

            label = f"[Fragment {i + 1}]  Source: {doc_name}"
            if page:
                label += f" (Page {page})"
            context_parts.append(f"{label}\n{doc.page_content}")

            if doc_name not in source_details:
                source_details[doc_name] = []
            source_details[doc_name].append({
                "chunk_index": i + 1,
                "page": page,
                "snippet": (
                    doc.page_content[:200] + "..."
                    if len(doc.page_content) > 200
                    else doc.page_content
                ),
            })

        sep = "\n\n" + "=" * 60 + "\n\n"
        context = sep.join(context_parts)
        context = "\n\n" + context + "\n\n" + "=" * 60

        # ── 3) Build user message (history + table_context + text context + question) ──
        history_text = ""
        if conversation_id:
            history = conversation_memory.get_history(conversation_id, max_turns=6)
            if history:
                history_parts = []
                for turn in history:
                    role_label = "User" if turn.role == "user" else "Assistant"
                    history_parts.append(f"{role_label}: {turn.content}")
                history_text = (
                    "## Conversation History\n\n"
                    + "\n\n".join(history_parts)
                    + "\n\n" + "=" * 60 + "\n\n"
                )

        # ── Comparison mode: re-group context by document ────────────
        if qtype == "comparison" and len(source_details) >= 2:
            doc_groups = __import__('collections').OrderedDict()
            for i, doc in enumerate(final_docs):
                doc_name = doc.metadata.get("document_name",
                                             doc.metadata.get("source", "Unknown"))
                page = str(doc.metadata.get("page") or "")
                if doc_name not in doc_groups:
                    doc_groups[doc_name] = []
                snippet = doc.page_content
                if page:
                    doc_groups[doc_name].append(f"[Page {page}]\n{snippet}")
                else:
                    doc_groups[doc_name].append(snippet)

            grouped_parts = []
            for doc_name, snippets in doc_groups.items():
                grouped_parts.append(
                    f"### Document: {doc_name}\n\n" + "\n\n".join(snippets)
                )
            comparison_context = "\n\n".join(grouped_parts)
            context = (
                "\n\n"
                "Note: The user is asking a cross-document comparison question. "
                "Below are passages grouped by source document.\n\n"
                + comparison_context
                + "\n\n"
            )
            context_label = "## Document Passages (grouped by document)\n"
            comparison_instruction = (
                "\n\n## Comparison Instruction\n\n"
                "This is a comparison question. Extract relevant data from each document "
                "and present a structured comparison. Use a table format where:\n"
                "- Each row = a comparison dimension (e.g. revenue, profit margin, penalty clause)\n"
                "- Each column = a document\n"
                "Then give a summary conclusion highlighting key differences and commonalities."
            )
        else:
            context_label = f"## Document Fragments (total {len(final_docs)})\n"
            comparison_instruction = ""

        user_message = (
            f"{history_text}"
            f"{table_context}"
            f"{context_label}"
            f"{context}\n\n"
            f"## Current Question\n\n"
            f"{question}"
            f"{comparison_instruction}"
        )

        # ── 4) LLM generation ──
        answer = self.llm_chat(get_domain_prompt(_domain), user_message)
        answer = self._format_answer(answer)

        convo_id = conversation_id or str(uuid.uuid4())
        conversation_memory.add_turn(
            convo_id,
            ConversationTurn(role="user", content=question),
        )
        conversation_memory.add_turn(
            convo_id,
            ConversationTurn(
                role="assistant",
                content=answer,
                sources=[s for s in source_details],
            ),
        )

        retrieve_method = (
            "hybrid + rerank" if use_hybrid and use_rerank
            else "hybrid" if use_hybrid
            else "vector + rerank" if use_rerank
            else "vector"
        )
        if qtype == "comparison" and len(source_details) >= 2:
            retrieve_method += " + comparison"
        elif table_context:
            retrieve_method += " + tables"

        result = {
            "question": question,
            "conversation_id": convo_id,
            "answer": answer,
            "sources": list(source_details.keys()),
            "source_details": [
                {"source": s, "chunks": d}
                for s, d in source_details.items()
            ],
            "source_count": len(final_docs),
            "retrieval_method": retrieve_method,
            "kb_id": kb_id,
        }

        return result

